import docx
import logging
from pathlib import Path

def create_word_document(pages_data: list[tuple[int, str]], output_path: str | Path):
    """
    Создает Word документ из списка кортежей (номер_страницы, текст_в_markdown).
    Каждый элемент списка будет отделен в документе.
    """
    logging.info(f"Создание Word документа: {output_path}")
    doc = docx.Document()
    doc.add_heading('Дневник', level=0)

    for i, (page_num, page_text) in enumerate(pages_data):
        # Добавляем разделитель страниц, кроме как перед первой страницей
        if i > 0:
            doc.add_page_break()
        
        doc.add_heading(f'Страница {page_num}', level=1) # Используем реальный номер страницы

        paragraphs = page_text.strip().split('\n\n')
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue

            if para.startswith('# '):
                doc.add_heading(para.lstrip('# ').strip(), level=2)
            elif para.startswith('## '):
                doc.add_heading(para.lstrip('## ').strip(), level=3)
            elif para.startswith('> '):
                p = doc.add_paragraph(para.lstrip('> ').strip())
                p.style = 'Quote'
            else:
                cleaned_para = para.replace('\n', ' ').strip()
                doc.add_paragraph(cleaned_para)

    try:
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        logging.info(f"Документ успешно сохранен: {output_path}")
    except Exception as e:
        logging.error(f"Не удалось сохранить Word документ: {e}")
