import os
import re
import base64
import requests
import time
from docx import Document
from docx.shared import Pt, Inches
from dotenv import load_dotenv

# --- КОНФИГУРАЦИЯ ---

# Загружаем переменные окружения из .env файла
load_dotenv()

# Получаем конфигурацию из переменных окружения
CONFIG = {
    "API_KEY": os.getenv("YC_API_KEY"),
    "FOLDER_ID": os.getenv("YC_FOLDER_ID"),
    "IMAGE_DIR": "scans_for_test",
    "OUTPUT_DOCX_NAME": "diary.docx",
    "GPT_MODEL_URI_TEMPLATE": "gpt://{folder_id}/yandexgpt/latest",
}

# --- ПРОМПТ ДЛЯ LLM ---
# Это самый важный компонент для сохранения аутентичности текста
SYSTEM_PROMPT = """
Ты — высокоточный ассистент-архивариус. Твоя задача — обработать сырой текст, полученный после автоматического распознавания (OCR) старого рукописного дневника. Этот текст имеет огромную историческую и личную ценность. Твоя работа требует предельной аккуратности.

Главная цель: отформатировать текст в Markdown, сохранив его 100% аутентичность.

Твои правила:
1.  **НЕ ИЗМЕНЯЙ СЛОВА И ФОРМУЛИРОВКИ.** Категорически запрещено перефразировать, подбирать синонимы, изменять порядок слов или "улучшать" стиль автора. Если в тексте написано "пошли мы с имярек гуляти", ты должен оставить "пошли мы с имярек гуляти", а не исправлять на "мы с имярек пошли гулять".
2.  **НЕ ИСПРАВЛЯЙ ГРАММАТИКУ И ОРФОГРАФИЮ.** Текст написан в СССР. Возможны устаревшие нормы, диалектизмы, авторские сокращения или просто ошибки. Ты должен их сохранить. Твоя задача — не быть корректором.
3.  **НЕ ДОДУМЫВАЙ.** Если слово распознано неверно или нечитаемо (например, "пришлoсь" или "пр#шлось"), не пытайся его угадать. Оставь как есть. Пользователь исправит это вручную.
4.  **ТВОЯ ЕДИНСТВЕННАЯ ЗАДАЧА — СТРУКТУРИРОВАНИЕ:**
    *   **Восстанови абзацы.** В сыром тексте все может быть слитно. Разделяй текст на логические абзацы, основываясь на смысле.
    *   **Определи заголовки.** Если строка явно выглядит как заголовок (например, дата "15 мая 1965 года" или название главы), отформатируй её как заголовок Markdown (например, `## 15 мая 1965 года`). Используй заголовки 2-го уровня (`##`).
    *   **Оформи цитаты.** Если в тексте есть цитаты (часто на немецком или английском), оформи их как цитаты Markdown (начав строку с `> `).
    *   **Объединяй разорванные слова.** OCR часто разрывает слова (например, "сло во"). Ты должен их аккуратно соединить ("слово"). Это единственное допустимое "исправление".
5.  **ВЫВОД — ТОЛЬКО MARKDOWN.** Не добавляй никаких своих комментариев, предисловий или заключений. Только отформатированный текст.

Пример:
Входной сырой текст: "15мая1965г. сегодня был а хорошая погода ходили в парк с аней она рассказала мне цитату Goethe Wer immer strebend sich bemüht, den können wir erlösen. потом мы пошли до дому."

Твой идеальный результат:
## 15 мая 1965 г.

Сегодня была хорошая погода. Ходили в парк с Аней.

Она рассказала мне цитату:
> Wer immer strebend sich bemüht, den können wir erlösen.

Потом мы пошли до дому.
"""

def check_config():
    """Проверяет наличие необходимых ключей в конфигурации."""
    if not CONFIG["API_KEY"] or not CONFIG["FOLDER_ID"]:
        print("!!! ОШИБКА: Не найдены YC_API_KEY или YC_FOLDER_ID.")
        print("!!! Пожалуйста, создайте файл .env и заполните его по инструкции.")
        return False
    return True

def natural_sort_key(s: str) -> list:
    """Ключ для "естественной" сортировки строк типа '10.jpg' > '2.jpg'."""
    return [int(text) if text.isdigit() else text.lower() for text in re.split(r'(\d+)', s)]

def encode_image_to_base64(filepath: str) -> str:
    """Кодирует файл изображения в строку Base64."""
    with open(filepath, "rb") as image_file:
        return base64.b64encode(image_file.read()).decode('utf-8')

def get_raw_text_from_ocr(base64_content: str) -> str | None:
    """Отправляет изображение в Yandex Vision OCR и возвращает сырой текст."""
    url = "https://ocr.api.cloud.yandex.net/ocr/v1/recognizeText"
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Api-Key {CONFIG['API_KEY']}",
        "x-folder-id": CONFIG["FOLDER_ID"],
    }
    body = {
        "mimeType": "JPEG",
        "languageCodes": ["ru", "en", "de"],
        "model": "handwritten",
        "content": base64_content
    }

    try:
        response = requests.post(url, headers=headers, json=body, timeout=120)
        if response.status_code == 200:
            result = response.json()
            full_text = result.get('result', {}).get('textAnnotation', {}).get('fullText', '')
            return full_text
        else:
            print(f"  [Ошибка OCR] Сервер вернул код: {response.status_code}. Ответ: {response.text}")
            return None
    except requests.exceptions.RequestException as e:
        print(f"  [Ошибка OCR] Не удалось подключиться к серверу: {e}")
        return None

def format_text_with_gpt(raw_text: str) -> str | None:
    """Отправляет сырой текст в YandexGPT для форматирования."""
    url = "https://llm.api.cloud.yandex.net/foundationModels/v1/completion"
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Api-Key {CONFIG['API_KEY']}",
    }
    body = {
        "modelUri": CONFIG["GPT_MODEL_URI_TEMPLATE"].format(folder_id=CONFIG["FOLDER_ID"]),
        "completionOptions": {
            "stream": False,
            "temperature": 0.1,  # Низкая температура для предсказуемости
            "maxTokens": "4000"
        },
        "messages": [
            {"role": "system", "text": SYSTEM_PROMPT},
            {"role": "user", "text": raw_text}
        ]
    }

    try:
        response = requests.post(url, headers=headers, json=body, timeout=180)
        if response.status_code == 200:
            result = response.json()
            formatted_text = result['result']['alternatives'][0]['message']['text']
            return formatted_text
        else:
            print(f"  [Ошибка GPT] Сервер вернул код: {response.status_code}. Ответ: {response.text}")
            return None
    except requests.exceptions.RequestException as e:
        print(f"  [Ошибка GPT] Не удалось подключиться к серверу: {e}")
        return None

def add_markdown_to_document(doc: Document, markdown_text: str):
    """Парсит Markdown и добавляет форматированный текст в Word документ."""
    for line in markdown_text.split('\n'):
        line = line.strip()
        if not line:
            continue
        
        if line.startswith('## '):
            doc.add_heading(line.lstrip('## ').strip(), level=2)
        elif line.startswith('# '):
            doc.add_heading(line.lstrip('# ').strip(), level=1)
        elif line.startswith('> '):
            p = doc.add_paragraph(style='Intense Quote')
            p.add_run(line.lstrip('> ').strip())
        else:
            doc.add_paragraph(line)

def main():
    """Основная функция скрипта."""
    if not check_config():
        return

    try:
        image_files = [f for f in os.listdir(CONFIG["IMAGE_DIR"]) if f.lower().endswith('.jpg')]
        image_files.sort(key=natural_sort_key)
    except FileNotFoundError:
        print(f"!!! ОШИБКА: Папка '{CONFIG['IMAGE_DIR']}' не найдена. Создайте ее и положите туда файлы.")
        return
        
    if not image_files:
        print(f"!!! ОШИБКА: В папке '{CONFIG['IMAGE_DIR']}' не найдено jpg файлов.")
        return

    print(f"Найдено {len(image_files)} страниц. Начинаю обработку...")

    doc = Document()
    
    total_files = len(image_files)
    start_time = time.time()

    for i, filename in enumerate(image_files, 1):
        page_number_match = re.search(r'(\d+)', filename)
        page_label = page_number_match.group(1) if page_number_match else filename

        print(f"[{i}/{total_files}] Обработка страницы: {filename}")
        
        doc.add_heading(f'Страница {page_label}', level=1)
        filepath = os.path.join(CONFIG["IMAGE_DIR"], filename)
        
        print("  -> Шаг 1: Распознавание рукописного текста (OCR)...")
        base64_content = encode_image_to_base64(filepath)
        raw_text = get_raw_text_from_ocr(base64_content)
        
        if not raw_text or not raw_text.strip():
            print("  [Предупреждение] Текст на странице не найден или пуст.")
            doc.add_paragraph("[На этой странице текст не найден]", style='Intense Quote')
            if i < total_files:
                doc.add_page_break()
            continue

        print(f"    [DEBUG]   Сырой текст:\n{raw_text}")

        print(f"  -> Шаг 2: Форматирование текста с помощью YandexGPT ({len(raw_text)} симв.)...")
        formatted_text = format_text_with_gpt(raw_text)

        print(f"    [DEBUG]   Отформатированный текст:\n{formatted_text}")

        if formatted_text:
            print("  -> Шаг 3: Добавление в Word документ...")
            add_markdown_to_document(doc, formatted_text)
        else:
            print("  [Предупреждение] Не удалось отформатировать текст. Вставляю сырой результат OCR.")
            doc.add_paragraph(raw_text)
        
        if i < total_files:
            doc.add_page_break()

    print("-" * 40)
    print("💾 Сохранение итогового файла...")
    doc.save(CONFIG["OUTPUT_DOCX_NAME"])
    
    end_time = time.time()
    total_time = end_time - start_time
    print("-" * 40)
    print("🎉 Обработка завершена!")
    print(f"Итоговый файл сохранен как: {CONFIG['OUTPUT_DOCX_NAME']}")
    print(f"Затраченное время: {total_time:.2f} секунд ({total_time/60:.2f} минут).")


if __name__ == '__main__':
    main()
